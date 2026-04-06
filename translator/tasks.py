"""
translator/tasks.py — translate DOCX preserving structure.
Uses argostranslate directly (offline, no internet needed).
"""
import os, io
from celery import shared_task
from django.utils import timezone


@shared_task(bind=True)
def translate_docx_task(self, job_id):
    from translator.models import TranslationJob
    from translator.utils import translate_long_text
    from django.core.files.base import ContentFile
    from docx import Document
    import uuid

    job = TranslationJob.objects.get(id=job_id)
    try:
        job.status = 'processing'
        job.save(update_fields=['status'])

        src = job.source_lang
        tgt = job.target_lang
        doc = Document(job.original_file.path)
        total_chars = 0

        # ── Translate paragraphs ───────────────────────────────
        for para in doc.paragraphs:
            original = para.text.strip()
            if not original:
                continue

            translated  = translate_long_text(original, source=src, target=tgt)
            total_chars += len(original)

            if para.runs:
                # Preserve formatting from first run
                first       = para.runs[0]
                bold        = first.bold
                italic      = first.italic
                font_name   = first.font.name
                font_size   = first.font.size
                try:
                    font_color = first.font.color.rgb if (
                        first.font.color and first.font.color.type) else None
                except Exception:
                    font_color = None

                # Clear all runs, write translated text into first
                for run in para.runs:
                    run.text = ''
                first.text   = translated
                first.bold   = bold
                first.italic = italic
                if font_name: first.font.name = font_name
                if font_size: first.font.size = font_size
                if font_color:
                    try: first.font.color.rgb = font_color
                    except Exception: pass
            else:
                para.add_run(translated)

        # ── Translate table cells ──────────────────────────────
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original = para.text.strip()
                        if not original:
                            continue
                        translated   = translate_long_text(original, source=src, target=tgt)
                        total_chars += len(original)
                        if para.runs:
                            for run in para.runs:
                                run.text = ''
                            para.runs[0].text = translated
                        else:
                            para.add_run(translated)

        # ── Save output ────────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        out_bytes = buf.read()

        stem        = os.path.splitext(job.original_name)[0]
        result_name = f"{stem}_{tgt.upper()}.docx"

        job.result_file.save(result_name, ContentFile(out_bytes), save=False)
        job.result_size  = len(out_bytes)
        job.char_count   = total_chars
        job.status       = 'done'
        job.save(update_fields=['result_file', 'result_size', 'char_count', 'status'])

    except Exception as e:
        job.status        = 'failed'
        job.error_message = str(e)
        job.save(update_fields=['status', 'error_message'])
        raise


@shared_task(bind=True)
def translate_pdf_task(self, job_id):
    from translator.models import TranslationJob
    from translator.utils import translate_long_text
    from editor.utils import extract_text_blocks, detect_pdf_type
    from django.core.files.base import ContentFile
    import fitz
    import os
    from collections import defaultdict

    job = TranslationJob.objects.get(id=job_id)
    try:
        job.status = 'processing'
        job.save(update_fields=['status'])

        doc = fitz.open(job.original_file.path)
        pdf_type = detect_pdf_type(doc)
        blocks = extract_text_blocks(doc)
        total_chars = 0

        # Group blocks by page
        page_map = defaultdict(list)
        for b in blocks:
            page_map[b['page']].append(b)

        for p_idx in range(len(doc)):
            page      = doc[p_idx]
            page_rect = page.rect
            page_blocks = page_map.get(p_idx, [])
            
            # Pass 1: Wipe original text via redaction
            for b in page_blocks:
                rect = fitz.Rect(b['x'], b['y'], b['x']+b['w'], b['y']+b['h'])
                page.add_redact_annot(rect, fill=(1,1,1))
            page.apply_redactions()

            # Pass 2: Insert translations in the exact same coordinates
            for b in page_blocks:
                original = b.get('text', '').strip()
                if not original: continue

                translated = translate_long_text(original, source=job.source_lang, target=job.target_lang)
                total_chars += len(original)

                x, y, w, h = b['x'], b['y'], b['w'], b['h']
                font_size = b.get('font_size', 12.0)
                font_name = _safe_font_fitz(b.get('font_name', ''))
                
                c = b.get('color', [0, 0, 0])
                if c == [255, 255, 255]: c = [0, 0, 0] # Avoid white on white
                color_f = (c[0]/255.0, c[1]/255.0, c[2]/255.0)

                # ── Key fix: Expand insertion rect to prevent overflow ──
                # Translated text is usually longer. We allow it to grow downwards.
                line_height = font_size * 1.3
                chars_per_line = max(1, int(w / (font_size * 0.5)))
                num_lines = max(1, -(-len(translated) // chars_per_line))
                needed_h = max(h, num_lines * line_height)
                
                insert_rect = fitz.Rect(x, y, min(page_rect.width, x + w), min(page_rect.height, y + needed_h))

                # Try inserting
                rc = page.insert_textbox(
                    insert_rect, 
                    translated,
                    fontsize=font_size,
                    fontname=font_name,
                    color=color_f,
                    align=0
                )

                # If still fails, try with smaller font
                if rc < 0:
                    rc2 = page.insert_textbox(
                        insert_rect,
                        translated,
                        fontsize=max(6.0, font_size * 0.7),
                        fontname=font_name,
                        color=color_f,
                        align=0
                    )
                    # Last resort: Single line insert_text (no wrap, but guaranteed to show)
                    if rc2 < 0:
                        page.insert_text(
                            fitz.Point(x, y + font_size),
                            translated[:100], # Truncate to prevent extreme overflow
                            fontsize=max(6.0, font_size * 0.6),
                            fontname=font_name,
                            color=color_f
                        )

        out_bytes = doc.tobytes(garbage=4, deflate=True, clean=True)
        doc.close()

        stem = os.path.splitext(job.original_name)[0]
        result_name = f"{stem}_{job.target_lang.upper()}.pdf"

        job.result_file.save(result_name, ContentFile(out_bytes), save=False)
        job.result_size = len(out_bytes)
        job.char_count = total_chars
        job.status = 'done'
        job.save(update_fields=['result_file', 'result_size', 'char_count', 'status'])

    except Exception as e:
        job.status = 'failed'
        job.error_message = str(e)
        job.save(update_fields=['status', 'error_message'])
        raise

def _safe_font_fitz(name):
    n = (name or '').lower()
    if any(k in n for k in ('helvetica','arial','sans')):
        if 'bold' in n and ('oblique' in n or 'italic' in n): return 'hebo'
        if 'oblique' in n or 'italic' in n: return 'heoi'
        return 'helv'
    if any(k in n for k in ('times','roman','serif')):
        if 'bold' in n and 'italic' in n: return 'tibi'
        if 'bold' in n:   return 'tibo'
        if 'italic' in n: return 'tiit'
        return 'tiro'
    if any(k in n for k in ('courier','mono','code')):
        if 'bold' in n and ('oblique' in n or 'italic' in n): return 'cobo'
        if 'oblique' in n or 'italic' in n: return 'coit'
        return 'cour'
    return 'helv'

@shared_task
def cleanup_translation_jobs():
    from translator.models import TranslationJob
    expired = TranslationJob.objects.filter(expires_at__lt=timezone.now())
    count = 0
    for job in expired:
        for f in [job.original_file, job.result_file]:
            if f:
                try:
                    if os.path.exists(f.path): os.remove(f.path)
                except Exception: pass
        job.delete()
        count += 1
    return f"Cleaned {count} expired translation jobs"
